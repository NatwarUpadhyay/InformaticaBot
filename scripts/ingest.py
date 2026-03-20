"""Ingestion script: reads docs, chunks, embeds, and stores in sqlite-vec"""
import os
import sys
import sqlite3
import numpy as np
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from sentence_transformers import SentenceTransformer
from src.config import DB_PATH

DOCS_DIR = "data/docs"
CHUNK_SIZE = 300
OVERLAP = 50
EMBED_MODEL = "all-MiniLM-L6-v2"


def extract_text_from_file(filepath):
    """Extract text from various file formats"""
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext == '.pdf':
            try:
                import PyPDF2
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text += page.extract_text()
                return text
            except ImportError:
                print(f"  Warning: PyPDF2 not installed. Install with: pip install PyPDF2")
                return ""
        
        elif ext in ['.docx', '.doc']:
            try:
                from docx import Document
                if ext == '.docx':
                    doc = Document(filepath)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text
                else:
                    # For .doc files, try using python-docx if available
                    doc = Document(filepath)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text
            except ImportError:
                print(f"  Warning: python-docx not installed. Install with: pip install python-docx")
                return ""
        
        elif ext in ['.md', '.txt']:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            print(f"  Warning: Unsupported file format: {ext}")
            return ""
    
    except Exception as e:
        print(f"  Error reading {filename}: {str(e)}")
        return ""


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


def read_docs():
    """Read all document files from data/docs/ and Sample Inputs/"""
    docs = []
    
    # Directories to scan
    doc_dirs = [Path(DOCS_DIR), Path("Sample Inputs")]
    
    for docs_path in doc_dirs:
        if not docs_path.exists():
            if docs_path.name == DOCS_DIR:
                os.makedirs(docs_path)
            continue
        
        # Support multiple file types
        supported_extensions = ['*.md', '*.txt', '*.pdf', '*.docx', '*.doc']
        
        for ext in supported_extensions:
            for p in sorted(docs_path.glob(ext)):
                text = extract_text_from_file(str(p))
                if text:  # Only add if text extraction was successful
                    # Use relative path for source
                    rel_path = f"{docs_path.name}/{p.name}"
                    docs.append({"id": rel_path, "text": text})
    
    return docs


def ingest():
    """Main ingestion function"""
    print(f"Loading embedding model: {EMBED_MODEL}")
    model = SentenceTransformer(EMBED_MODEL)

    print(f"Reading documents from {DOCS_DIR}...")
    docs = read_docs()
    
    if not docs:
        print(f"No documents found in {DOCS_DIR}. Please add some .md or .txt files.")
        return

    print(f"Found {len(docs)} document(s). Chunking...")
    
    all_chunks = []
    all_sources = []
    
    for doc in docs:
        chunks = chunk_text(doc["text"])
        for chunk in chunks:
            all_chunks.append(chunk)
            all_sources.append(doc["id"])
        print(f"  {doc['id']}: {len(chunks)} chunks")

    print(f"Total chunks: {len(all_chunks)}")
    print(f"Computing embeddings...")
    embeddings = model.encode(all_chunks, show_progress_bar=True)
    embedding_dim = embeddings.shape[1]
    
    print(f"Embedding dimension: {embedding_dim}")
    print(f"Creating SQLite database at {DB_PATH}...")
    
    # Create connection and load sqlite-vec extension
    conn = sqlite3.connect(DB_PATH)
    
    try:
        # Load the sqlite-vec extension
        conn.enable_load_extension(True)
        conn.load_extension("vec0")
        conn.enable_load_extension(False)
    except Exception as e:
        print(f"Warning: Could not load vec0 extension: {e}")
        print("Using basic SQLite storage instead (without vector similarity search)")
    
    # Drop old tables
    conn.execute("DROP TABLE IF EXISTS chunks")
    conn.execute("DROP TABLE IF EXISTS vec_chunks")
    
    # Create chunks table
    conn.execute("""
        CREATE TABLE chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            chunk_text TEXT NOT NULL,
            source TEXT NOT NULL
        )
    """)
    
    # Try to create vector table; fall back to regular table if vec0 not available
    try:
        conn.execute(f"""
            CREATE VIRTUAL TABLE vec_chunks USING vec0(
                chunk_id INTEGER PRIMARY KEY,
                embedding float[{embedding_dim}]
            )
        """)
        use_vector_table = True
    except Exception as e:
        print(f"Warning: Could not create vec0 table: {e}")
        print("Creating basic embedding storage table instead")
        conn.execute("""
            CREATE TABLE vec_chunks (
                chunk_id INTEGER PRIMARY KEY,
                embedding BLOB
            )
        """)
        use_vector_table = False
    
    # Insert data
    print("Inserting chunks and embeddings...")
    for chunk, source, embedding in zip(all_chunks, all_sources, embeddings):
        cursor = conn.execute("INSERT INTO chunks (chunk_text, source) VALUES (?, ?)", (chunk, source))
        chunk_id = cursor.lastrowid
        
        if use_vector_table:
            # Store embedding as binary blob
            embedding_bytes = embedding.astype(np.float32).tobytes()
            conn.execute("INSERT INTO vec_chunks (rowid, embedding) VALUES (?, ?)", (chunk_id, embedding_bytes))
        else:
            # Store as binary blob
            embedding_bytes = embedding.astype(np.float32).tobytes()
            conn.execute("INSERT INTO vec_chunks (chunk_id, embedding) VALUES (?, ?)", (chunk_id, embedding_bytes))
    
    conn.commit()
    conn.close()
    print(f"✓ Successfully ingested {len(all_chunks)} chunks into {DB_PATH}")


if __name__ == "__main__":
    ingest()


if __name__ == "__main__":
    ingest()
